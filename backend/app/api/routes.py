from fastapi import APIRouter, HTTPException, BackgroundTasks
from app.models.schemas import MaterialRequest, GenerationResponse
from app.agents.orchestrator import IntelligentOrchestrator
from app.models.config import MaterialConfig
from app.tools.storage import save_to_history, get_history
import logging
import asyncio
import traceback
from typing import List

router = APIRouter()
logger = logging.getLogger("API")

@router.post("/generate")
async def generate_math_material(request: MaterialRequest, background_tasks: BackgroundTasks):
    """
    Starter generering i bakgrunnen for å unngå timeout.
    """
    logger.info(f"Mottatt forespørsel: {request.emne} ({request.klassetrinn})")
    
    def run_generation_sync():
        """Synkron funksjon som kjører i bakgrunnen."""
        try:
            logger.info(f"Bakgrunnsjobb starter for: {request.emne}")
            
            config = MaterialConfig(
                klassetrinn=request.klassetrinn,
                emne=request.emne,
                kompetansemaal=request.kompetansemaal,
                differentiation=request.differentiation,
                include_answer_key=request.include_answer_key,
                document_format=request.document_format
            )
            
            orchestrator = IntelligentOrchestrator()
            crew = orchestrator.create_dynamic_crew(config)
            
            logger.info("Crew opprettet, starter kickoff...")
            result = crew.kickoff()
            logger.info("Crew kickoff ferdig!")
            
            final_code = result.raw if hasattr(result, 'raw') else str(result)
            
            # Rens koden for vanlige AI-feil
            from app.core.sanitizer import sanitize_typst_code
            if config.document_format.value == "typst":
                final_code = sanitize_typst_code(final_code)
            else:
                # For LaTeX, bare fjern markdown fences
                import re
                final_code = re.sub(r'```(?:typst|latex)?\n?', '', final_code)
                final_code = final_code.replace('```', '').strip()
            
            logger.info(f"Kode generert og renset ({len(final_code)} tegn), starter kompilering...")
            
            # Kompiler til PDF
            from app.core.compiler import DocumentCompiler, TypstTemplates
            compiler = DocumentCompiler()
            
            # Fjern AI-generert preamble hvis den finnes
            lines = final_code.split('\n')
            content_lines = []
            skip_preamble = True
            for line in lines:
                if skip_preamble and line.strip().startswith('#set'):
                    continue
                skip_preamble = False
                content_lines.append(line)
            
            content = '\n'.join(content_lines).strip()
            
            # Legg til vår preamble
            preamble = TypstTemplates.worksheet_header(
                title=config.emne,
                grade=config.klassetrinn,
                topic=config.emne
            )
            final_code = preamble + "\n" + content

            worksheet_pdf = None
            figures = []
            
            # HYBRID MODE: Generer figurer hvis nødvendig
            is_hybrid = config.document_format.value == "hybrid"
            if is_hybrid:
                logger.info("Hybrid-modus aktivert, genererer figurer...")
                try:
                    figures = orchestrator.generate_figures(config)
                    logger.info(f"Generert {len(figures)} figurer")
                except Exception as e:
                    logger.warning(f"Figurgenering feilet: {e}")
                    figures = []
            
            # Kompiler PDF
            import subprocess
            import tempfile
            from pathlib import Path
            import base64
            
            try:
                with tempfile.TemporaryDirectory() as tmpdir:
                    tmpdir_path = Path(tmpdir)
                    typ_file = tmpdir_path / "document.typ"
                    pdf_file = tmpdir_path / "document.pdf"
                    
                    # Opprett figur-mappe hvis hybrid
                    if figures:
                        fig_dir = tmpdir_path / "figurer"
                        fig_dir.mkdir(exist_ok=True)
                        
                        # Kompiler TikZ-figurer til PNG
                        for fig in figures:
                            logger.info(f"Kompilerer figur: {fig['id']}")
                            try:
                                # Kjør async kompilering synkront
                                loop = asyncio.new_event_loop()
                                asyncio.set_event_loop(loop)
                                try:
                                    fig_result = loop.run_until_complete(
                                        compiler.compile_latex_figure_to_png(fig['latex'])
                                    )
                                    if fig_result.success and fig_result.png_bytes:
                                        png_path = fig_dir / f"{fig['id']}.png"
                                        png_path.write_bytes(fig_result.png_bytes)
                                        logger.info(f"Figur {fig['id']} lagret som PNG")
                                    else:
                                        logger.warning(f"Figur {fig['id']} feilet: {fig_result.log}")
                                finally:
                                    loop.close()
                            except Exception as e:
                                logger.warning(f"Kunne ikke kompilere figur {fig['id']}: {e}")
                    
                    typ_file.write_text(final_code, encoding="utf-8")
                    logger.info(f"Typst-fil skrevet: {len(final_code)} tegn")
                    
                    result = subprocess.run(
                        ["typst", "compile", str(typ_file), str(pdf_file)],
                        capture_output=True,
                        timeout=90,
                        cwd=tmpdir
                    )
                    
                    if pdf_file.exists():
                        pdf_bytes = pdf_file.read_bytes()
                        worksheet_pdf = base64.b64encode(pdf_bytes).decode("utf-8")
                        logger.info(f"PDF kompilert! Størrelse: {len(pdf_bytes)} bytes")
                    else:
                        logger.error(f"Typst feilet. stdout: {result.stdout.decode()}")
                        logger.error(f"Typst feilet. stderr: {result.stderr.decode()}")
            except FileNotFoundError:
                logger.error("Typst er ikke installert på serveren!")
            except subprocess.TimeoutExpired:
                logger.error("Typst-kompilering timet ut")
            except Exception as e:
                logger.error(f"Kompileringsfeil: {str(e)}")
            
            save_to_history(config, worksheet_pdf if worksheet_pdf else "", None, final_code)
            logger.info(f"Bakgrunnsjobb FERDIG for: {request.emne}")
            
        except Exception as e:
            logger.error(f"Bakgrunnsgenerering feilet: {str(e)}")
            logger.error(traceback.format_exc())
            # Lagre feilet forsøk med feilmelding
            try:
                save_to_history(
                    MaterialConfig(
                        klassetrinn=request.klassetrinn,
                        emne=f"[FEILET] {request.emne}",
                        kompetansemaal=request.kompetansemaal
                    ),
                    "",
                    None,
                    f"% Generering feilet: {str(e)}"
                )
            except:
                pass

    background_tasks.add_task(run_generation_sync)
    
    return {
        "success": True, 
        "message": "Generering startet i bakgrunnen. Sjekk Oppgavebanken om 1-2 minutter.",
        "status": "processing"
    }

@router.get("/history")
async def fetch_history(limit: int = 10):
    """Henter genereringshistorikken."""
    try:
        return get_history(limit)
    except Exception as e:
        logger.error(f"Kunne ikke hente historikk: {e}")
        return []

@router.get("/health")
async def health_check():
    return {"status": "healthy", "version": "v3.0-pro-templates"}

@router.post("/export/word")
async def export_to_word(request: MaterialRequest):
    """Eksporterer generert innhold til Word-format."""
    try:
        from app.tools.word_exporter import is_word_export_available, latex_to_word
        import tempfile
        import base64
        
        if not is_word_export_available():
            raise HTTPException(status_code=503, detail="Word-eksport er ikke tilgjengelig")
        
        # Hent siste genererte dokument for dette emnet
        history = get_history(limit=10)
        matching = [h for h in history if h.get('emne') == request.emne]
        
        if not matching:
            raise HTTPException(status_code=404, detail="Ingen generert innhold funnet for dette emnet")
        
        source_code = matching[0].get('source_code', '')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            # For nå, eksporter kildekoden som tekst i Word
            # TODO: Implementer Typst-til-Word konvertering
            from docx import Document
            doc = Document()
            doc.add_heading(f"{request.emne} - {request.klassetrinn}", 0)
            
            # Legg til innhold
            for line in source_code.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            doc.save(tmp.name)
            
            with open(tmp.name, 'rb') as f:
                word_bytes = f.read()
        
        return {
            "success": True,
            "word_b64": base64.b64encode(word_bytes).decode("utf-8"),
            "filename": f"{request.emne}_{request.klassetrinn}.docx"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Word-eksport feilet: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/test-typst")
async def test_typst():
    """Tester om Typst fungerer på serveren."""
    import subprocess
    import tempfile
    from pathlib import Path
    
    test_code = """#set text(size: 12pt)
= Test
Dette er en test av Typst-kompilering.
$ x^2 + y^2 = z^2 $
"""
    
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            typ_file = Path(tmpdir) / "test.typ"
            pdf_file = Path(tmpdir) / "test.pdf"
            typ_file.write_text(test_code)
            
            result = subprocess.run(
                ["typst", "compile", str(typ_file), str(pdf_file)],
                capture_output=True,
                timeout=30
            )
            
            if pdf_file.exists():
                return {
                    "status": "ok",
                    "message": "Typst fungerer!",
                    "pdf_size": pdf_file.stat().st_size
                }
            else:
                return {
                    "status": "error",
                    "stdout": result.stdout.decode(),
                    "stderr": result.stderr.decode()
                }
    except FileNotFoundError:
        return {"status": "error", "message": "Typst er ikke installert"}
    except Exception as e:
        return {"status": "error", "message": str(e)}
